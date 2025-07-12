# -*- coding: utf-8 -*-
import io
from datetime import datetime
from pathlib import Path  # Importação importante para caminhos robustos
from fastapi import FastAPI, Request, Form
from fastapi.responses import StreamingResponse
from fastapi.templating import Jinja2Templates
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mangum import Mangum

# Nossas classes e motor
from modelos import Paciente, Laudo
from motor_laudos import GeradorDeLaudos

# --- CONFIGURAÇÃO ROBUSTA DE CAMINHOS ---
# Cria um caminho absoluto para o diretório onde este script está.
# Isso garante que os arquivos (json, jpg) sejam encontrados, não importa de onde o script seja executado.
BASE_DIR = Path(__file__).resolve().parent

# --- CONFIGURAÇÃO DA APLICAÇÃO WEB ---
# O Netlify serve a pasta 'publish' ('templates') na raiz do site.
# No entanto, a função serverless roda a partir de 'functions' ('api').
# Para o Jinja2 encontrar os templates, precisamos ser explícitos.
# Como o publish dir é "templates", o caminho relativo a partir do build root é correto.
# Nota: Esta configuração de template é para rodar localmente. No Netlify, ele servirá o HTML estaticamente.
# O ponto crucial é que a função serverless precisa encontrar seus próprios arquivos (JSONs, JPG).
templates = Jinja2Templates(directory=str(BASE_DIR.parent / "templates"))

app = FastAPI(title="API Gerador de Laudos")

MAPA_DE_EXAMES = {
    # Agora usamos BASE_DIR para criar o caminho completo para cada arquivo.
    "joelho": BASE_DIR / "joelho.json",
    "abdominal": BASE_DIR / "abdominal.json"
}

@app.get("/")
def home(request: Request):
    """Esta rota serve o formulário HTML para testes locais."""
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/gerar-laudo")
async def gerar_laudo_endpoint(
    tipo_exame: str = Form(...),
    nome_paciente: str = Form(...),
    data_nascimento: str = Form(...),
    medico_solicitante: str = Form(...),
    descricao_medico: str = Form(...)
):
    caminho_json = MAPA_DE_EXAMES.get(tipo_exame)
    if not caminho_json:
        return {"error": "Tipo de exame não suportado"}

    motor = GeradorDeLaudos(caminho_base_conhecimento=caminho_json)
    
    info_exame = motor.get_info_exame()
    paciente_obj = Paciente(nome=nome_paciente, data_nascimento=data_nascimento)
    lista_de_achados = motor.processar_descricao(descricao_medico)

    laudo_obj = Laudo(
        paciente=paciente_obj,
        medico_solicitante=medico_solicitante,
        lado_exame="", 
        template_tecnica=info_exame["template_tecnica"]
    )
    for achado in lista_de_achados:
        laudo_obj.adicionar_achado(achado)

    document = Document()
    
    # Adiciona o logo usando o caminho robusto
    caminho_logo = BASE_DIR / 'artrofisio.jpg'
    p_logo = document.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_logo.add_run()
    run.add_picture(str(caminho_logo), width=Inches(2.5))

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # O resto da geração do documento permanece igual
    document.add_heading(info_exame.get("nome_exame", "LAUDO DE ULTRASSONOGRAFIA").upper(), level=1)
    document.add_paragraph(f"Paciente: {laudo_obj.paciente.nome} (Idade: {laudo_obj.paciente.get_idade()} anos)")
    document.add_paragraph(f"Médico Solicitante: Dr(a). {laudo_obj.medico_solicitante}")
    document.add_paragraph(f"Data do Exame: {datetime.now().strftime('%d de %B de %Y')}")
    document.add_heading('TÉCNICA:', level=2)
    document.add_paragraph(laudo_obj.template_tecnica.format(lado=""))
    document.add_heading('ACHADOS:', level=2)
    texto_achados = "\n".join([f"- {achado.texto_formal}" for achado in laudo_obj.achados])
    document.add_paragraph(texto_achados if laudo_obj.achados else "- Estruturas avaliadas sem alterações ecográficas significativas.")
    document.add_heading('IMPRESSÃO DIAGNÓSTICA:', level=2)
    document.add_paragraph(laudo_obj.gerar_impressao_diagnostica())

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    
    return StreamingResponse(
        file_stream,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename="laudo_{paciente_obj.nome.replace(" ", "_")}.docx"'}
    )

# "Cola" para o Netlify (AWS Lambda) chamar nossa aplicação FastAPI
handler = Mangum(app)